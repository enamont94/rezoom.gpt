"""
PDF and DOCX Parser Utilities
Handles parsing of resume files in various formats
"""

import fitz  # PyMuPDF
from docx import Document
import io
import re
from typing import Dict, Any, List, Optional
import logging

logger = logging.getLogger(__name__)

class PDFParser:
    """PDF parsing utilities using PyMuPDF"""
    
    def __init__(self):
        self.supported_formats = ['.pdf']
    
    def parse_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """
        Parse PDF file and extract text content
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            Dictionary with parsed content
        """
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            text_content = ""
            metadata = {}
            
            # Extract text from all pages
            for page_num in range(len(doc)):
                page = doc[page_num]
                text_content += page.get_text()
            
            # Extract metadata
            metadata = doc.metadata
            
            # Extract images (optional)
            images = self.extract_images(doc)
            
            doc.close()
            
            return {
                'text': text_content,
                'metadata': metadata,
                'page_count': len(doc),
                'images': images,
                'format': 'pdf'
            }
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise Exception(f"PDF parsing failed: {str(e)}")
    
    def extract_images(self, doc) -> List[Dict[str, Any]]:
        """
        Extract images from PDF document
        
        Args:
            doc: PyMuPDF document object
            
        Returns:
            List of image information
        """
        images = []
        
        try:
            for page_num in range(len(doc)):
                page = doc[page_num]
                image_list = page.get_images()
                
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        images.append({
                            'page': page_num + 1,
                            'index': img_index,
                            'width': pix.width,
                            'height': pix.height,
                            'colorspace': pix.colorspace.name if pix.colorspace else 'unknown'
                        })
                    
                    pix = None
            
        except Exception as e:
            logger.warning(f"Error extracting images: {str(e)}")
        
        return images
    
    def extract_tables(self, file_content: bytes) -> List[Dict[str, Any]]:
        """
        Extract tables from PDF
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            List of extracted tables
        """
        tables = []
        
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_tables = page.find_tables()
                
                for table in page_tables:
                    table_data = table.extract()
                    if table_data:
                        tables.append({
                            'page': page_num + 1,
                            'data': table_data,
                            'bbox': table.bbox
                        })
            
            doc.close()
            
        except Exception as e:
            logger.warning(f"Error extracting tables: {str(e)}")
        
        return tables

class DOCXParser:
    """DOCX parsing utilities using python-docx"""
    
    def __init__(self):
        self.supported_formats = ['.docx', '.doc']
    
    def parse_docx(self, file_content: bytes) -> Dict[str, Any]:
        """
        Parse DOCX file and extract content
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Dictionary with parsed content
        """
        try:
            doc = Document(io.BytesIO(file_content))
            text_content = ""
            paragraphs = []
            tables = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
                    text_content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Extract metadata
            metadata = {
                'title': doc.core_properties.title,
                'author': doc.core_properties.author,
                'created': doc.core_properties.created,
                'modified': doc.core_properties.modified,
                'subject': doc.core_properties.subject,
                'keywords': doc.core_properties.keywords
            }
            
            return {
                'text': text_content,
                'paragraphs': paragraphs,
                'tables': tables,
                'metadata': metadata,
                'format': 'docx'
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise Exception(f"DOCX parsing failed: {str(e)}")
    
    def extract_styles(self, file_content: bytes) -> Dict[str, Any]:
        """
        Extract style information from DOCX
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Dictionary with style information
        """
        try:
            doc = Document(io.BytesIO(file_content))
            styles = {}
            
            # Extract paragraph styles
            paragraph_styles = []
            for paragraph in doc.paragraphs:
                if paragraph.style.name:
                    paragraph_styles.append({
                        'text': paragraph.text[:50] + "..." if len(paragraph.text) > 50 else paragraph.text,
                        'style': paragraph.style.name,
                        'font_size': paragraph.style.font.size.pt if paragraph.style.font.size else None,
                        'bold': paragraph.style.font.bold,
                        'italic': paragraph.style.font.italic
                    })
            
            styles['paragraphs'] = paragraph_styles
            
            return styles
            
        except Exception as e:
            logger.warning(f"Error extracting styles: {str(e)}")
            return {}

class ResumeParser:
    """Main resume parser that handles multiple formats"""
    
    def __init__(self):
        self.pdf_parser = PDFParser()
        self.docx_parser = DOCXParser()
    
    def parse_resume(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse resume file based on file extension
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Dictionary with parsed content
        """
        try:
            # Get file extension
            file_extension = filename.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                return self.pdf_parser.parse_pdf(file_content)
            elif file_extension in ['docx', 'doc']:
                return self.docx_parser.parse_docx(file_content)
            else:
                raise Exception(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error parsing resume: {str(e)}")
            raise Exception(f"Resume parsing failed: {str(e)}")
    
    def extract_structured_data(self, parsed_content: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract structured data from parsed content
        
        Args:
            parsed_content: Parsed content from parser
            
        Returns:
            Dictionary with structured resume data
        """
        text = parsed_content.get('text', '')
        
        if not text:
            return {}
        
        # Extract contact information
        contact_info = self.extract_contact_info(text)
        
        # Extract sections
        sections = self.extract_sections(text)
        
        # Extract skills
        skills = self.extract_skills(text)
        
        # Extract experience
        experience = self.extract_experience(text)
        
        # Extract education
        education = self.extract_education(text)
        
        return {
            'contact_info': contact_info,
            'sections': sections,
            'skills': skills,
            'experience': experience,
            'education': education,
            'raw_text': text,
            'metadata': parsed_content.get('metadata', {})
        }
    
    def extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information from text"""
        contact_info = {
            'name': '',
            'email': '',
            'phone': '',
            'address': '',
            'linkedin': '',
            'website': ''
        }
        
        # Extract name (usually first line or first few words)
        lines = text.split('\n')
        if lines:
            first_line = lines[0].strip()
            if len(first_line.split()) <= 3:  # Likely a name
                contact_info['name'] = first_line
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info['email'] = email_match.group()
        
        # Extract phone
        phone_pattern = r'(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact_info['phone'] = phone_match.group()
        
        # Extract LinkedIn
        linkedin_pattern = r'linkedin\.com/in/([a-zA-Z0-9-]+)'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact_info['linkedin'] = linkedin_match.group(1)
        
        # Extract website
        website_pattern = r'(?:https?://)?(?:www\.)?([a-zA-Z0-9-]+\.[a-zA-Z]{2,})'
        website_match = re.search(website_pattern, text)
        if website_match:
            contact_info['website'] = website_match.group(1)
        
        return contact_info
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """Extract sections from resume text"""
        sections = {}
        
        # Common section headers
        section_patterns = {
            'summary': ['summary', 'objective', 'profile', 'about'],
            'experience': ['experience', 'work history', 'employment', 'career'],
            'education': ['education', 'academic', 'qualifications'],
            'skills': ['skills', 'technical skills', 'competencies'],
            'certifications': ['certifications', 'certificates'],
            'projects': ['projects', 'portfolio'],
            'languages': ['languages', 'language skills']
        }
        
        lines = text.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            line_clean = line.strip()
            if not line_clean:
                continue
            
            # Check if line is a section header
            is_section_header = False
            for section_name, keywords in section_patterns.items():
                if any(keyword in line_clean.lower() for keyword in keywords):
                    # Save previous section
                    if current_section and current_content:
                        sections[current_section] = '\n'.join(current_content)
                    
                    # Start new section
                    current_section = section_name
                    current_content = []
                    is_section_header = True
                    break
            
            if not is_section_header and current_section:
                current_content.append(line_clean)
        
        # Save the last section
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text"""
        # Common technical skills
        technical_skills = [
            'javascript', 'python', 'java', 'react', 'node.js', 'sql', 'aws', 'docker',
            'kubernetes', 'git', 'agile', 'scrum', 'machine learning', 'data analysis',
            'html', 'css', 'typescript', 'angular', 'vue', 'mongodb', 'postgresql'
        ]
        
        text_lower = text.lower()
        found_skills = []
        
        for skill in technical_skills:
            if skill in text_lower:
                found_skills.append(skill.title())
        
        return list(set(found_skills))
    
    def extract_experience(self, text: str) -> List[Dict[str, str]]:
        """Extract work experience from resume text"""
        experience = []
        
        # Look for experience section
        experience_section = ""
        lines = text.split('\n')
        in_experience = False
        
        for line in lines:
            line_clean = line.strip()
            if any(keyword in line_clean.lower() for keyword in ['experience', 'work history', 'employment']):
                in_experience = True
                continue
            elif in_experience and line_clean:
                if any(keyword in line_clean.lower() for keyword in ['education', 'skills', 'certifications']):
                    break
                experience_section += line_clean + "\n"
        
        # Simple parsing of experience (this could be enhanced)
        if experience_section:
            # Look for job titles and companies
            lines = experience_section.split('\n')
            current_job = {}
            
            for line in lines:
                line_clean = line.strip()
                if not line_clean:
                    continue
                
                # Look for job title patterns
                if any(word in line_clean.lower() for word in ['engineer', 'developer', 'manager', 'analyst', 'specialist']):
                    if current_job:
                        experience.append(current_job)
                    current_job = {
                        'title': line_clean,
                        'company': '',
                        'duration': '',
                        'description': ''
                    }
                elif current_job and not current_job['company']:
                    current_job['company'] = line_clean
                elif current_job and any(char.isdigit() for char in line_clean):
                    current_job['duration'] = line_clean
                elif current_job:
                    current_job['description'] += line_clean + " "
            
            if current_job:
                experience.append(current_job)
        
        return experience
    
    def extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education information from resume text"""
        education = []
        
        # Look for education section
        education_section = ""
        lines = text.split('\n')
        in_education = False
        
        for line in lines:
            line_clean = line.strip()
            if any(keyword in line_clean.lower() for keyword in ['education', 'academic', 'qualifications']):
                in_education = True
                continue
            elif in_education and line_clean:
                if any(keyword in line_clean.lower() for keyword in ['experience', 'skills', 'certifications']):
                    break
                education_section += line_clean + "\n"
        
        # Simple parsing of education
        if education_section:
            lines = education_section.split('\n')
            for line in lines:
                line_clean = line.strip()
                if line_clean and any(word in line_clean.lower() for word in ['bachelor', 'master', 'phd', 'degree', 'university', 'college']):
                    education.append({
                        'degree': line_clean,
                        'institution': '',
                        'year': '',
                        'gpa': ''
                    })
        
        return education
